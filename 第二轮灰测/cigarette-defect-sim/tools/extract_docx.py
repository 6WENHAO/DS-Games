import docx
p = r"E:\姜啸雷 基于改进的 YOLOv11 的可移动 上传终稿 6.11\姜啸雷 基于改进的 YOLOv11 的可移动 上传终稿 6.11\姜啸雷 基于改进的 YOLOv11 的可移动 上传终稿 6.11.docx"
d = docx.Document(p)
out = []
for para in d.paragraphs:
    t = para.text.strip()
    if t:
        out.append((para.style.name, t))
txt = "\n".join(f"[{s}] {t}" for s, t in out)
base = r"C:\Users\Owen\dsh-work\cigdefect\tools"
open(base + r"\thesis_text.txt", "w", encoding="utf-8").write(txt)
tb = []
for i, tbl in enumerate(d.tables):
    tb.append(f"=== TABLE {i} ===")
    for row in tbl.rows:
        tb.append(" | ".join(c.text.strip().replace("\n", " ") for c in row.cells))
open(base + r"\thesis_tables.txt", "w", encoding="utf-8").write("\n".join(tb))
print("paras", len(out), "chars", len(txt), "tables", len(d.tables))
